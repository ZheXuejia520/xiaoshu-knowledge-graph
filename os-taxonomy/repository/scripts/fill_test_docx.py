"""
模拟教师填写模板，生成一份测试 docx，用于验证导入解析流程。
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

def fill_cell(cell, text):
    """填充单元格内容"""
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def fill_table(table, values):
    """按行填充表格"""
    for i, value in enumerate(values):
        if i < len(table.rows):
            fill_cell(table.rows[i].cells[1], value)

def main():
    template_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "知识点录入模板.docx")
    doc = Document(template_path)
    tables = doc.tables

    # 表1：基本信息
    fill_table(tables[0], [
        "cn-sx-619",           # ID
        "有理数的概念",         # 名称
        "数学",                # 学科
        "6",                   # 年级
        "数与代数",            # 领域
        "1",                   # 排序号
        "六年级数学知识点导图"  # 来源
    ])

    # 表2：基础标签
    fill_table(tables[1], [
        "第三学段（5-6年级）",   # 学段
        "第一单元 · 有理数",     # 章节
        "第1课时",              # 课时
        "理解有理数的意义，能用数轴上的点表示有理数"  # 课标核心要求
    ])

    # 表3：图谱逻辑链路
    fill_table(tables[2], [
        "cn-sx-114,cn-sx-204,cn-sx-309",  # 前置
        "cn-sx-620,cn-sx-621,cn-sx-622"   # 后置
    ])

    # 表4：核心目标
    fill_table(tables[3], [
        "理解有理数的概念，掌握有理数的分类，能区分整数、分数、正数、负数",
        "培养分类讨论和抽象概括的能力",
        "发展数感和符号意识，体会数系扩充的必要性"
    ])

    # 表5：知识内容
    fill_table(tables[4], [
        "有理数是可以表示为两个整数之比的数，即形如 a/b（b≠0）的数",
        "有理数包括整数和分数\n正有理数大于0，负有理数小于0\n0既不是正有理数也不是负有理数",
        "有理数 = {a/b | a∈Z, b∈Z, b≠0}\n正有理数 + 负有理数 = 0（互为相反数）"
    ])

    # 表6：知识描述与示例
    fill_table(tables[5], [
        "理解有理数的概念，能正确区分有理数和无理数，掌握有理数的分类方法",
        "例1：判断下列数哪些是有理数：3, -5, 0.5, -2/3, π\n例2：将下列有理数分类：-3, 0, 1/2, -0.8, 5",
        "记住有理数就是能写成分数形式的数\n整数也是有理数，因为整数可以写成 整数/1",
        "把π当成有理数\n认为0既不是正数也不是负数，所以不是有理数",
        "能正确判断一个数是否为有理数\n能对有理数进行正确分类"
    ])

    # 表7：五步法
    fill_table(tables[6], [
        "判断下列各数中哪些是有理数：3.14, -√2, 0, 1/3, π",
        "有理数 = 能表示为两个整数之比的数；√2和π不能表示为两个整数之比",
        "1. 3.14 = 314/100，是有理数\n2. -√2 ≈ -1.414...，无限不循环，不是有理数\n3. 0 = 0/1，是有理数\n4. 1/3 本身就是分数形式，是有理数\n5. π ≈ 3.14159...，无限不循环，不是有理数",
        "有理数：3.14, 0, 1/3；非有理数：-√2, π",
        "有理数的本质是'可比数'，即可以用两个整数的比来表示。整数和有限小数、无限循环小数都是有理数，无限不循环小数不是有理数。"
    ])

    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "test_import", "有理数的概念.docx")
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"测试文件已生成：{output_path}")

if __name__ == "__main__":
    main()