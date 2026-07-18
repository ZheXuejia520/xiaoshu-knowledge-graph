"""
生成「数学知识图谱知识点录入表」Word 模板
运行：python scripts/generate_template.py
输出：知识点录入模板.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color):
    """设置单元格底色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def add_section(doc, title, fields, description=""):
    """添加一个章节，包含标题 + 说明 + 表格"""
    doc.add_heading(title, level=2)
    if description:
        p = doc.add_paragraph(description)
        p.style = doc.styles['Normal']
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)

    table = doc.add_table(rows=len(fields), cols=2, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, hint) in enumerate(fields):
        # 标签列
        cell_label = table.cell(i, 0)
        cell_label.text = ""
        p = cell_label.paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.bold = True
        set_cell_shading(cell_label, "E8EDF2")
        # 设置标签列宽度
        cell_label.width = Cm(4)

        # 内容列
        cell_value = table.cell(i, 1)
        cell_value.text = ""
        p = cell_value.paragraphs[0]
        run = p.add_run(hint)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(150, 150, 150)
        cell_value.width = Cm(12)

    doc.add_paragraph("")  # 间距
    return table

def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading('数学知识图谱 · 知识点录入表', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('请按照以下模板填写知识点信息，灰色文字为填写提示，填写时请替换为实际内容。\n'
                     '带 * 的为必填项。列表类字段请用换行（Enter）分隔多条内容。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph("")

    # ====== 一、基本信息 ======
    add_section(doc, "一、基本信息", [
        ("知识点 ID *", "如：cn-sx-xxx（数学用 cn-sx，语文 cn-yw，英语 cn-en，科学 cn-sc）"),
        ("知识点名称 *", "如：数数与一一对应"),
        ("学科 *", "数学 / 语文 / 英语 / 科学"),
        ("年级 *", "1 / 2 / 3 / 4 / 5 / 6"),
        ("领域 *", "如：数与代数、图形与几何、统计与概率、综合实践"),
        ("排序号", "同年级同领域内的排序，如：1"),
        ("来源", "如：一年级代数知识点导图"),
    ])

    # ====== 二、基础标签 ======
    add_section(doc, "二、基础标签（菜单栏 01）", [
        ("学段", "如：第一学段（1-2年级）。留空则自动从年级推断"),
        ("章节 *", "如：第一单元 · 准备课"),
        ("课时", "如：第1课时"),
        ("课标核心要求", "如：能根据给定的标准对事物进行分类..."),
    ])

    # ====== 三、图谱逻辑链路 ======
    add_section(doc, "三、图谱逻辑链路（菜单栏 02）", [
        ("前置知识点 ID", "多个用英文逗号分隔，如：cn-sx-102,cn-sx-103。无则留空"),
        ("后置知识点 ID", "多个用英文逗号分隔。无则留空"),
    ], description="注意：此处仅填写本知识点之外的前置/后置依赖关系，数据层面的依赖关系在 dependencies 中统一维护。")

    # ====== 四、核心目标 ======
    add_section(doc, "四、核心目标（菜单栏 03）", [
        ("知识目标 *", "如：掌握逐个数、分组数等多种计数方法..."),
        ("能力目标", "如：培养有序观察和分类计数的能力..."),
        ("素养目标", "如：发展数感和符号意识..."),
    ])

    # ====== 五、知识内容 ======
    add_section(doc, "五、知识内容（菜单栏 04）", [
        ("定义", "如：数数是指按照一定的顺序，将物体与自然数一一对应..."),
        ("性质", "每行一条性质，如：\n按顺序数：从小到大依次计数\n分组数：可以两个、五个一组地数"),
        ("公式", "每行一条公式，如：\na + b × c = a + (b × c)\n(a + b) × c ≠ a + b × c"),
    ])

    # ====== 六、知识描述与示例 ======
    add_section(doc, "六、知识描述与示例", [
        ("描述 *", "一句话概括这个知识点讲什么"),
        ("举例", "每行一个例子"),
        ("学习技巧", "每行一条技巧"),
        ("常见错误", "每行一条常见错误"),
        ("掌握表现", "每行一条掌握证据"),
    ])

    # ====== 七、知识练习（五步法） ======
    add_section(doc, "七、知识练习 · 五步法（菜单栏 05）", [
        ("审题", "题目内容"),
        ("建模", "建立数学模型的过程"),
        ("分步", "分步骤解答"),
        ("检验", "如何验证答案"),
        ("反思", "总结反思"),
    ], description="填写一个典型例题，按照五步法展开。")

    # 页脚说明
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("—— 填写完毕后，请将本文件发回，系统将自动解析并更新知识图谱 ——")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True

    # 保存
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "知识点录入模板.docx")
    doc.save(output_path)
    print(f"模板已生成：{output_path}")

if __name__ == "__main__":
    main()